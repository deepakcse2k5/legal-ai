import os
import fitz  # PyMuPDF for PDF extraction
from docx import Document  # python-docx for DOCX extraction & saving
from langchain_community.llms import Ollama
from src.config import MODEL_NAME, UPLOAD_FOLDER, OUTPUT_FOLDER

# Configuration

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Initialize LLM
llm = Ollama(model=MODEL_NAME)


# Utility Functions
def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()


def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()


def sanitize_text(text):
    """Replace Unicode characters with ASCII equivalents."""
    replacements = {
        "\u2014": "--",  # Em dash
        "\u2013": "-",  # En dash
        "\u2018": "'", "\u2019": "'",  # Curly single quotes
        "\u201C": '"', "\u201D": '"',  # Curly double quotes
        "\u2026": "...",  # Ellipsis
        "\u00A0": " "  # Non-breaking space
    }
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text


def generate_complaint_report(reference_text, case_details):
    """Generate a legally formatted complaint report based on provided details."""
    prompt = f"""
    You are a legal expert specializing in drafting complaint reports. Your task is to generate a new complaint report **strictly following the format, tone, and structure** of the provided reference complaint report.

    ### **Reference Complaint Report**
    {reference_text}

    ---

    **Instructions:**  
    - Use the reference report as a strict template. Maintain the same section order, formatting, legal language, and writing style.
    - Ensure logical coherence while integrating the details of the new case.
    - **Maintain formal and legally appropriate phrasing.**

    ### **New Complaint Report Details**
    {case_details}
    """

    response = llm.invoke(prompt)
    return sanitize_text(response.strip())


def save_report_as_docx(report_text, filename):
    """Save the complaint report as a DOCX file."""
    doc = Document()
    doc.add_heading("Complaint Report", level=1)
    for line in report_text.split("\n"):
        doc.add_paragraph(line)
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(filepath)
    print(f"✅ Report successfully saved at: {filepath}")  # Debugging
    return filepath


# Example Usage
def process_complaint(use_pdf=True, reference_pdf_path="", reference_docx_path="", case_details={}):
    """Process a complaint report based on PDF/DOCX input and save output."""
    reference_text = extract_text_from_pdf(reference_pdf_path) if use_pdf else extract_text_from_docx(
        reference_docx_path)
    report_text = generate_complaint_report(reference_text, case_details)
    filename = f"Complaint_Report_{case_details['petitioner'].replace(' ', '_')}.docx"
    return save_report_as_docx(report_text, filename)
